"""
This module defines functions for handling input and output of data.

"""

from datetime import datetime
import os
import traceback

import pandas as pd
import docx

def create_generator(input_, batch_size=10):
    """
    Generator function for yielding values in batches based on the type
    of file passed.

    Parameters
    ----------
    input : str or iterable
        Path to file containing list of identifiers
    batch_size : int, optional
        Number of items in each batch, default 10
    
    Yields
    ------
    list
        List of size `batch_size` of the passed iterable of objects

    """

    if os.path.isfile(input_):

        if input_.endswith('.docx'):
            doc = docx.Document(input_)
            number_of_lines = len(doc.paragraphs)

            for count in range(0, number_of_lines, batch_size):
                yield [paragraph.text.encode('utf-8','ignore').strip() 
                for paragraph in doc.paragraphs[count:min(count + batch_size, number_of_lines)]]
            
        # Add options to read other types of files 
        #
        # elif input_.endswith('.pdf'):
        #     identifier_generator =  None # Process .pdf file
        # elif input_.endswith('.txt'):
        #     identifier_generator = None # Process .txt file
        # elif input_.endswith('.xlsx'):
        #     identifier_generator = None # Process .xlsx file
        
    else:

        for identifier_list in [[input_]]:
            yield identifier_list

def export_final_df(final_df):
    """
    Apply conditional formatting to the passed DataFrame and write to an Excel file.

    Parameters
    ----------
    final_df : pandas.DataFrame
        Pandas dataframe object with at least four columns and `TRUE` 
        and `FALSE` values

    """

    try:
        writer = pd.ExcelWriter(f'ICRIS Download Status - {datetime.today().strftime("%c")}.xlsx', engine = 'xlsxwriter')

        final_df.to_excel(writer)

        workbook = writer.book
        worksheet = writer.sheets['Sheet1']

        format1 = workbook.add_format({'bg_color': '#FFC7CE'})
        format2 = workbook.add_format({'bg_color': '#C6EFCE'})

        worksheet.conditional_format(f'D2:D{str(len(final_df.index) + 1)}', {
                                            'type': 'cell',
                                            'criteria': '==',
                                            'value': 'TRUE',
                                            'format': format2
                                            })

        worksheet.conditional_format(f'D2:D{str(len(final_df.index) + 1)}', {'type': 'cell',
                                            'criteria': '==',
                                            'value': 'FALSE',
                                            'format': format1
                                            })                                

        writer.save()

        print("\n\n\t\t****Excel file written successfully***\n\n")

    except Exception:
        tb = traceback.format_exc()
        print(tb)
        print(f"\n\n\t\t****Excel file could not be written****\n\n")
